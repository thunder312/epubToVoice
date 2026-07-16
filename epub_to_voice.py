#!/usr/bin/env python3
"""
epub_to_voice.py – EPUB / PDF / TXT to Audio Converter
Uses edge-tts (Microsoft Edge TTS, free, no API key needed) for high-quality neural voices.

Structure detection (headings, paragraphs) uses SSML <break> tags (Option 1).
Where SSML is not possible, spoken placeholder sentences are used (Option 3) with a warning.
Table of contents sections are automatically detected and skipped.

Usage:
    python epub_to_voice.py book.epub
    python epub_to_voice.py book.pdf -o ./output -v de-DE-KatjaNeural -r "+10%"
    python epub_to_voice.py book.txt
    python epub_to_voice.py --list-voices | grep de-DE
"""

import asyncio
import argparse
import html as _html
import json
import re
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path

import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import edge_tts


# ---------------------------------------------------------------------------
# Defaults
# ---------------------------------------------------------------------------
DEFAULT_VOICE  = "de-DE-ConradNeural"
DEFAULT_RATE   = "-10%"
DEFAULT_VOLUME = "+0%"

# A few recommended alternative voices (see --list-voices for the full,
# regularly-updated catalog of all Edge TTS languages/voices):
#   German female : de-DE-KatjaNeural, de-DE-AmalaNeural,
#                    de-DE-SeraphinaMultilingualNeural (expressive, multilingual)
#   German male   : de-DE-ConradNeural (default), de-DE-KillianNeural,
#                    de-DE-FlorianMultilingualNeural
#   English female: en-US-JennyNeural, en-US-AriaNeural, en-GB-SoniaNeural
#   English male  : en-US-GuyNeural, en-US-AndrewMultilingualNeural, en-GB-RyanNeural


# ---------------------------------------------------------------------------
# Segment types & tag sets
# ---------------------------------------------------------------------------
SEG_HEADING   = "heading"
SEG_PARAGRAPH = "paragraph"

KEEP_TAGS    = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "td", "th", "blockquote"}
HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
SKIP_TAGS    = {"nav", "script", "style", "head"}

# TOC chapter title detection (case-insensitive)
TOC_TITLE_RE = re.compile(
    r"^\s*(inhalts?verzeichnis|inhalt|contents?|table\s+of\s+contents?|toc|sommaire|índice)\s*$",
    re.IGNORECASE,
)

# ---------------------------------------------------------------------------
# Audio-level pause & prosody design (edge-tts XML-escapes all text, so
# in-band SSML <break>/<prosody> tags are never interpreted - see
# Communicate.__init__ in the edge_tts package, which always runs the text
# through xml.sax.saxutils.escape(). Real, exactly-timed pauses and
# heading/subheading differentiation are only achievable by synthesising
# segments individually with per-level pitch/rate and splicing in actual
# silent audio via ffmpeg - see plan_audio_pieces() / _synthesise_chunk_edge.)
# ---------------------------------------------------------------------------
_HEADING_PITCH_HZ  = {1: 15, 2: 8, 3: 3}     # extra pitch per heading level (Hz)
_HEADING_RATE_PCT  = {1: -6, 2: -3, 3: -1}   # extra (slower) rate per heading level (%)
_DEFAULT_HEADING_PITCH_HZ = 0
_DEFAULT_HEADING_RATE_PCT = 0

_PAUSE_BEFORE_HEADING_MS = {1: 900, 2: 650, 3: 450}   # silence before a heading
_PAUSE_AFTER_HEADING_MS  = {1: 500, 2: 350, 3: 250}   # silence after a heading, before body text
_DEFAULT_PAUSE_BEFORE_HEADING_MS = 400
_DEFAULT_PAUSE_AFTER_HEADING_MS  = 200
_PAUSE_PARAGRAPH_MS   = 220     # silence between two consecutive paragraphs
_PAUSE_SCENE_BREAK_MS = 1500    # silence for a scene-break marker ("***", "* * *", ...)

# Standalone symbol-only paragraph = scene break marker
SCENE_BREAK_RE = re.compile(r"^[\s*#•·×~=\-–—]{2,}$")


# ---------------------------------------------------------------------------
# Text / segment helpers
# ---------------------------------------------------------------------------

_DASH_RE = re.compile(r"[–—]")  # – (en-dash, U+2013) and — (em-dash, U+2014)


def replace_dashes_for_tts(text: str) -> str:
    """
    Replace en-dash (–) and em-dash (—) with a comma.

    Edge TTS reads these characters aloud literally as the word "Dash"
    instead of pausing, so they must be rewritten before synthesis.
    """
    if not text or not _DASH_RE.search(text):
        return text
    text = _DASH_RE.sub(",", text)
    text = re.sub(r"\s*,\s*", ", ", text)            # normalize spacing around commas
    text = re.sub(r"(,\s*){2,}", ", ", text)         # collapse repeated commas
    text = re.sub(r"^\s*,\s*", "", text)             # drop leading comma (dialogue dash)
    text = re.sub(r"\s*,\s*([.!?…])", r"\1", text)   # drop comma right before terminal punctuation
    return text.strip()


def sanitize_text(text: str) -> str:
    """Remove control characters and normalize whitespace."""
    text = text.replace("\ufffd", "")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


# ---------------------------------------------------------------------------
# Edge-TTS dialogue/pause preprocessing
# ---------------------------------------------------------------------------
# edge-tts synthesises each paragraph break in the extracted text as its own
# call (see plan_audio_pieces), which always gets a hard splice pause - even
# mid-dialogue or mid-sentence, where the source text never got a real
# terminal punctuation mark. A closing dialogue quote (") does NOT by itself
# count as a real sentence end below, since German dialogue lines almost
# always end in one and would otherwise never merge with the next line.
_EDGE_TTS_SENTENCE_END_RE   = re.compile(r"[.!?…—]\s*$")
_EDGE_TTS_DIALOGUE_COMMA_RE = re.compile(r"“,$", re.MULTILINE)  # “,  ->  “
_EDGE_TTS_CHAPTER_RE        = re.compile(r"^Kapitel\b")
_EDGE_TTS_SHORT_PARA_LEN    = 8   # paragraphs shorter than this always merge back


def merge_incomplete_segments(segments: list[dict]) -> list[dict]:
    """
    Merge consecutive SEG_PARAGRAPH segments where the preceding segment
    does not end in a real sentence boundary (. ! ? … —), or where the
    current segment is a very short fragment (edge-tts drops/glitches on
    mini-segments). Heading segments and scene-break markers ("***" etc.)
    are always kept as-is and act as merge barriers.
    This removes unwanted TTS pauses caused by mid-sentence/mid-dialogue
    paragraph breaks - e.g. EPUB markup that puts each dialogue line in its
    own <p> tag.
    """
    result: list[dict] = []
    for seg in segments:
        text = _EDGE_TTS_DIALOGUE_COMMA_RE.sub("“", seg["text"])
        mergeable = (
            seg["type"] == SEG_PARAGRAPH
            and not is_scene_break(text)
            and result
            and result[-1]["type"] == SEG_PARAGRAPH
            and not is_scene_break(result[-1]["text"])
            and (
                len(text) < _EDGE_TTS_SHORT_PARA_LEN
                or not _EDGE_TTS_SENTENCE_END_RE.search(result[-1]["text"])
            )
        )
        if mergeable:
            result[-1] = {
                "type": SEG_PARAGRAPH,
                "text": result[-1]["text"].rstrip() + " " + text.lstrip(),
            }
        else:
            new_seg = dict(seg)
            new_seg["text"] = text
            result.append(new_seg)
    return result


def prepare_text_for_edge_tts(text: str) -> str:
    """
    Preprocess extracted text before handing it to edge-tts so paragraph
    breaks only produce a spoken pause where the text actually ends a
    sentence or dialogue line - not mid-dialogue or mid-sentence.

    - Paragraphs not ending in . ! ? … — are joined to the next paragraph
      with a single space.
    - Paragraphs ending in one of those are kept as their own line (natural
      speaking pause).
    - Paragraphs shorter than 8 characters always merge into the previous
      one (edge-tts drops/glitches on mini-segments).
    - "Kapitel ..." headings get a blank line before and after.
    - Stylistic trailing commas after a closing dialogue quote are removed,
      e.g. „Nein, Mylord.",  ->  „Nein, Mylord."
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _EDGE_TTS_DIALOGUE_COMMA_RE.sub("“", text)

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    blocks: list[str] = []
    for para in paragraphs:
        is_barrier = bool(_EDGE_TTS_CHAPTER_RE.match(para)) or is_scene_break(para)
        if not blocks:
            blocks.append(para)
            continue

        prev = blocks[-1]
        prev_is_barrier = bool(_EDGE_TTS_CHAPTER_RE.match(prev)) or is_scene_break(prev)
        if is_barrier or prev_is_barrier:
            blocks.append(para)
        elif len(para) < _EDGE_TTS_SHORT_PARA_LEN or not _EDGE_TTS_SENTENCE_END_RE.search(prev):
            blocks[-1] = prev.rstrip() + " " + para
        else:
            blocks.append(para)

    # Blocks that keep their own "line break" become real paragraphs
    # (separated by a blank line), which downstream parsing (\n{2,}-split)
    # turns back into individual TTS segments with a natural pause between
    # them - as opposed to merged blocks, which share one segment/API call.
    out_parts: list[str] = [blocks[0]] if blocks else []
    for i in range(1, len(blocks)):
        is_heading      = bool(_EDGE_TTS_CHAPTER_RE.match(blocks[i]))
        prev_is_heading = bool(_EDGE_TTS_CHAPTER_RE.match(blocks[i - 1]))
        out_parts.append("\n\n\n" if (is_heading or prev_is_heading) else "\n\n")
        out_parts.append(blocks[i])

    return "".join(out_parts)


def is_toc_title(title: str) -> bool:
    """Return True if the title matches common TOC heading patterns."""
    return bool(TOC_TITLE_RE.match(title.strip()))


def is_toc_content(text: str) -> bool:
    """Heuristic: True if >40% of non-empty lines look like TOC entries (end with page numbers)."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if len(lines) < 5:
        return False
    toc_lines = sum(1 for l in lines if re.search(r"[\.\s]{3,}\d+\s*$", l))
    return toc_lines / len(lines) > 0.4


def segments_to_plain_text(segments: list[dict]) -> str:
    """Collapse segments to plain text (used for length checks and TOC detection)."""
    return " ".join(seg["text"] for seg in segments)


def _rate_to_piper_length_scale(rate: str) -> float:
    """Convert an Edge-TTS rate string like '+20%' or '-10%' to a Piper length_scale.

    Piper length_scale: 1.0 = normal, >1.0 = slower, <1.0 = faster.
    Edge rate: +N% = faster, -N% = slower → inverse mapping.
    """
    rate = rate.strip()
    if not rate or rate == "+0%":
        return 1.0
    try:
        sign = 1 if rate.startswith("+") else -1
        pct  = float(rate.lstrip("+-").rstrip("%"))
        # +20% faster in Edge → factor 1/1.20 in Piper
        factor = 1.0 + sign * pct / 100.0
        return round(1.0 / max(factor, 0.1), 3)
    except ValueError:
        return 1.0


def segments_to_tts_text(segments: list[dict]) -> str:
    """
    Build TTS-ready text with punctuation-based prosodic pauses (Option 1).

    NOTE: edge-tts XML-escapes the text before embedding it in SSML, so literal
    <break> tags would be read aloud. Instead we use punctuation-based prosodic
    pauses that the TTS engine interprets naturally:
      • Before a heading  → "…" (ellipsis, ~1 s pause)
      • After  a heading  → "." (period, short pause before body text)
      • Between paragraphs→ text already ends with punctuation; one space suffices
    """
    _SENTENCE_END = frozenset('.!?…\u201d\u00bb')   # ."!?…"»

    parts: list[str] = []
    for i, seg in enumerate(segments):
        text = replace_dashes_for_tts(seg["text"].strip())
        if not text:
            continue

        if seg["type"] == SEG_HEADING:
            if i > 0:
                parts.append(" \u2026 ")      # " … " – long prosodic pause before heading
            # Ensure heading ends with a sentence-boundary so TTS pauses after it
            if text[-1] not in _SENTENCE_END:
                text += "."
            parts.append(text)
            parts.append(" ")                 # brief gap before body text
        else:
            if i > 0:
                prev_is_heading = segments[i - 1]["type"] == SEG_HEADING
                if not prev_is_heading:
                    # Add a short pause between paragraphs via sentence boundary
                    if parts and parts[-1][-1] not in _SENTENCE_END:
                        parts.append(".")
                    parts.append(" ")
                else:
                    parts.append(" ")
            parts.append(text)

    return "".join(parts)


def segments_to_fallback_text(segments: list[dict]) -> str:
    """
    Option 3 fallback: spoken placeholder sentences at structural boundaries.
    Used when the primary synthesis fails at runtime.
    """
    parts: list[str] = []
    for i, seg in enumerate(segments):
        if i > 0:
            parts.append("Neues Kapitel." if seg["type"] == SEG_HEADING else "Weiter.")
        parts.append(replace_dashes_for_tts(seg["text"]))
    return " ".join(parts)


def chunk_segments(segments: list[dict], max_chars: int = 4900) -> list[list[dict]]:
    """
    Group segments into chunks whose combined text fits within max_chars.
    Splits within a segment only when a single segment exceeds max_chars.
    """
    chunks: list[list[dict]] = []
    current: list[dict] = []
    current_len = 0

    for seg in segments:
        seg_len = len(seg["text"])

        if seg_len > max_chars:
            # Flush current chunk first
            if current:
                chunks.append(current)
                current, current_len = [], 0
            # Hard-split the oversized segment at sentence boundaries
            sentences = re.split(r"(?<=[.!?])\s+", seg["text"])
            buf = ""

            def _split_piece(text: str) -> dict:
                piece = {"type": seg["type"], "text": text}
                if "level" in seg:
                    piece["level"] = seg["level"]
                return piece

            for sent in sentences:
                if len(buf) + len(sent) + 1 <= max_chars:
                    buf = (buf + " " + sent).strip()
                else:
                    if buf:
                        chunks.append([_split_piece(buf)])
                    while len(sent) > max_chars:
                        chunks.append([_split_piece(sent[:max_chars])])
                        sent = sent[max_chars:]
                    buf = sent
            if buf:
                chunks.append([_split_piece(buf)])

        elif current_len + seg_len + 1 > max_chars:
            chunks.append(current)
            current, current_len = [seg], seg_len

        else:
            current.append(seg)
            current_len += seg_len + 1

    if current:
        chunks.append(current)

    return chunks


def is_scene_break(text: str) -> bool:
    """True if text is a standalone scene-break marker ("***", "* * *", "###", ...)."""
    return bool(SCENE_BREAK_RE.match(text.strip()))


def _heading_prosody(level: int | None) -> tuple[str, int]:
    """Return (pitch string, rate-delta-percent) for a heading of the given level."""
    lvl = level or 1
    hz  = _HEADING_PITCH_HZ.get(lvl, _DEFAULT_HEADING_PITCH_HZ)
    pct = _HEADING_RATE_PCT.get(lvl, _DEFAULT_HEADING_RATE_PCT)
    return f"+{hz}Hz", pct


def _combine_rate(base_rate: str, delta_pct: int) -> str:
    """Apply a percentage-point delta on top of a '+N%'/'-N%' base rate string."""
    m = re.match(r"^([+-])(\d+)%$", base_rate)
    if not m:
        return base_rate
    value = int(m.group(2)) if m.group(1) == "+" else -int(m.group(2))
    value += delta_pct
    return f"{'+' if value >= 0 else '-'}{abs(value)}%"


def plan_audio_pieces(segments: list[dict], base_rate: str) -> list[dict]:
    """
    Turn a list of segments into an ordered list of audio "pieces" for
    per-segment synthesis + silence-splicing:
      {"kind": "tts", "text": ..., "pitch": ..., "rate": ...}
      {"kind": "silence", "ms": ...}

    Heading levels get progressively more pitch/rate emphasis and longer
    surrounding pauses (level 1 = strongest), so headings and sub-headings
    become audibly distinguishable. Scene-break markers ("***" etc.) become
    a single longer silence with no spoken text.
    """
    pieces: list[dict] = []
    prev_kind: str | None = None   # "heading" | "paragraph" | "scene_break"

    for seg in segments:
        text = replace_dashes_for_tts(seg["text"].strip())
        if not text:
            continue

        if seg["type"] == SEG_HEADING:
            level = seg.get("level")
            if prev_kind is not None:
                ms = _PAUSE_BEFORE_HEADING_MS.get(level or 1, _DEFAULT_PAUSE_BEFORE_HEADING_MS)
                pieces.append({"kind": "silence", "ms": ms})
            pitch, rate_delta = _heading_prosody(level)
            pieces.append({
                "kind": "tts", "text": text,
                "pitch": pitch, "rate": _combine_rate(base_rate, rate_delta),
            })
            ms = _PAUSE_AFTER_HEADING_MS.get(level or 1, _DEFAULT_PAUSE_AFTER_HEADING_MS)
            pieces.append({"kind": "silence", "ms": ms})
            prev_kind = "heading"

        elif is_scene_break(text):
            pieces.append({"kind": "silence", "ms": _PAUSE_SCENE_BREAK_MS})
            prev_kind = "scene_break"

        else:
            if prev_kind == "paragraph":
                pieces.append({"kind": "silence", "ms": _PAUSE_PARAGRAPH_MS})
            pieces.append({"kind": "tts", "text": text, "pitch": "+0Hz", "rate": base_rate})
            prev_kind = "paragraph"

    return pieces


# ---------------------------------------------------------------------------
# EPUB helpers
# ---------------------------------------------------------------------------

def html_to_segments(html_bytes: bytes) -> list[dict]:
    """Extract structured segments (headings + paragraphs) from an EPUB HTML chunk."""
    soup = BeautifulSoup(html_bytes, "html.parser")
    for tag in soup(SKIP_TAGS):
        tag.decompose()

    segments: list[dict] = []
    for element in soup.find_all(KEEP_TAGS):
        text = sanitize_text(element.get_text(" ", strip=True))
        if not text:
            continue
        is_hdg = element.name in HEADING_TAGS
        seg: dict = {"type": SEG_HEADING if is_hdg else SEG_PARAGRAPH, "text": text}
        if is_hdg:
            seg["level"] = int(element.name[1])   # h1..h6 → 1..6
        segments.append(seg)

    # Fallback: no recognised tags → treat full text as one paragraph
    if not segments:
        text = sanitize_text(soup.get_text(" ", strip=True))
        if text:
            segments.append({"type": SEG_PARAGRAPH, "text": text})

    return segments


# ---------------------------------------------------------------------------
# Custom exceptions
# ---------------------------------------------------------------------------

class TtsServerError(RuntimeError):
    """Raised when the TTS server is unreachable or returns a service error."""


# ---------------------------------------------------------------------------
# HTML conversion log
# ---------------------------------------------------------------------------

class HtmlLog:
    """
    Writes a colour-coded HTML log file to the output directory.
    Each entry gets an elapsed-time timestamp.
    Also mirrors every message to stdout.
    """

    _CSS = """
    body{font-family:'Consolas','Courier New',monospace;background:#1e1e2e;
         color:#cdd6f4;padding:28px;margin:0;line-height:1.55}
    h1{color:#89b4fa;margin:0 0 4px;font-size:1.25em}
    .meta{color:#6c7086;margin-bottom:22px;font-size:.88em}
    .meta b{color:#a6adc8;margin-right:18px}
    table{width:100%;border-collapse:collapse;font-size:.93em}
    tr:hover td{background:#24243e}
    td{padding:2px 6px;vertical-align:top}
    .ts{color:#45475a;white-space:nowrap;width:72px;font-size:.82em;padding-top:3px}
    .ok   {color:#a6e3a1}
    .error{color:#f38ba8;font-weight:bold}
    .warn {color:#fab387}
    .info {color:#cdd6f4}
    .dim  {color:#585b70}
    .sep  {color:#313244;padding:4px 0}
    .footer{margin-top:24px;padding-top:14px;border-top:1px solid #313244;
            color:#6c7086;font-size:.88em}
    .s-ok {color:#a6e3a1} .s-err{color:#f38ba8} .s-warn{color:#fab387}
    """

    def __init__(self, path: Path, title: str, meta: dict, enabled: bool = True) -> None:
        self._path    = path
        self._enabled = enabled
        self._f       = open(path, "w", encoding="utf-8") if enabled else None
        self._start   = datetime.now()
        self._counts: dict[str, int] = {"ok": 0, "error": 0, "warn": 0}
        if self._enabled:
            self._write_header(title, meta)

    # ------------------------------------------------------------------
    def _write_header(self, title: str, meta: dict) -> None:
        started = self._start.strftime("%d.%m.%Y %H:%M:%S")
        rows = "".join(
            f'<b>{_html.escape(k)}:</b>{_html.escape(str(v))} '
            for k, v in meta.items()
        )
        self._f.write(f"""<!DOCTYPE html>
<html lang="de">
<head>
  <meta charset="UTF-8">
  <title>Protokoll – {_html.escape(title)}</title>
  <style>{self._CSS}</style>
</head>
<body>
<h1>📖 {_html.escape(title)}</h1>
<div class="meta">{rows}<b>Start:</b>{started}</div>
<table>
""")
        self._f.flush()

    # ------------------------------------------------------------------
    def _elapsed(self) -> str:
        t   = int((datetime.now() - self._start).total_seconds())
        h, r = divmod(t, 3600)
        m, s = divmod(r, 60)
        return f"{h:02d}:{m:02d}:{s:02d}"

    # ------------------------------------------------------------------
    def log(self, msg: str, level: str = "info") -> None:
        """Write one log line (level: info | ok | warn | error | dim | sep). Always mirrored to stdout."""
        ts  = self._elapsed()
        if level == "sep":
            if self._enabled:
                self._f.write(f'<tr><td colspan="2" class="sep"><hr></td></tr>\n')
                self._f.flush()
            print()
            return

        if level in self._counts:
            self._counts[level] += 1

        if self._enabled:
            safe = _html.escape(msg)
            self._f.write(f'<tr><td class="ts">{ts}</td><td class="{level}">{safe}</td></tr>\n')
            self._f.flush()
        print(msg)

    # ------------------------------------------------------------------
    def close(self, total_chunks: int, produced: int) -> None:
        if not self._enabled:
            return
        end      = datetime.now()
        duration = str(end - self._start).split(".")[0]
        ok   = self._counts["ok"]
        err  = self._counts["error"]
        warn = self._counts["warn"]
        self._f.write(f"""</table>
<div class="footer">
  Splits gesamt: {total_chunks} &nbsp;|&nbsp;
  <span class="s-ok">✓ {ok} erfolgreich</span> &nbsp;|&nbsp;
  <span class="s-warn">⚠ {warn} Warnungen</span> &nbsp;|&nbsp;
  <span class="s-err">✗ {err} Fehler</span> &nbsp;|&nbsp;
  Dauer: {duration} &nbsp;|&nbsp;
  Abgeschlossen: {end.strftime('%d.%m.%Y %H:%M:%S')}
</div>
</body></html>
""")
        self._f.close()
        print(f"\n📋  Protokoll gespeichert: {self._path}")


# ---------------------------------------------------------------------------
# Core converter
# ---------------------------------------------------------------------------

class EpubConverter:
    def __init__(
        self,
        epub_path: str | Path,
        output_dir: str | Path | None = None,
        voice: str = DEFAULT_VOICE,
        rate: str = DEFAULT_RATE,
        volume: str = DEFAULT_VOLUME,
        skip_short: int = 60,
        max_chapters: int | None = None,
        merge: bool = False,
        start_page: int | None = None,      # PDF: first page to include (1-based)
        end_page: int | None = None,        # PDF: last page to include (1-based, inclusive)
        skip_chapters: set[int] | None = None,  # EPUB/TXT: doc-order indices to skip
        translate_to: str | None = None,        # BCP-47 language code to translate into (e.g. "de")
        resume: bool = False,                   # skip splits whose output file already exists
        tts_engine: str = "edge",               # "edge" or "piper"
        piper_voice: str = "de_DE-thorsten-high",
        normalize_volume: bool = True,          # loudness-normalize splits via ffmpeg (best effort)
        save_log: bool = False,                 # write the HTML Protokoll file to the output directory
        optimize_before_reading: bool = True,   # splice real pauses at paragraph/heading/scene breaks
    ):
        self.epub_path    = Path(epub_path)
        self.output_dir   = Path(output_dir) if output_dir else self.epub_path.parent
        self.voice         = voice
        self.rate          = rate
        self.volume        = volume
        self.skip_short    = skip_short
        self.max_chapters  = max_chapters
        self.merge         = merge
        self.start_page    = start_page
        self.end_page      = end_page
        self.skip_chapters = skip_chapters or set()
        self.translate_to  = translate_to or None
        self.resume        = resume
        self.tts_engine       = tts_engine
        self.piper_voice      = piper_voice
        self.normalize_volume = normalize_volume
        self.save_log         = save_log
        self.optimize_before_reading = optimize_before_reading
        self._piper_instance  = None   # lazy-loaded

    @property
    def _audio_ext(self) -> str:
        return ".mp3"  # both Edge and Piper output MP3

    # ------------------------------------------------------------------
    @staticmethod
    def _clean_title(raw: str) -> str:
        """Strip author prefix/suffix from a raw title string."""
        for sep in (" - ", " – ", ": "):
            if sep in raw:
                raw = raw.split(sep, 1)[-1].strip()
                break
        return re.sub(r"\s*\([^)]*\)\s*$", "", raw).strip()

    # ------------------------------------------------------------------
    def _load_chapters(self) -> tuple[list[dict], str]:
        """Dispatch to the right loader based on file extension."""
        ext = self.epub_path.suffix.lower()
        if ext == ".epub":
            return self._load_epub()
        if ext == ".pdf":
            return self._load_pdf()
        if ext == ".txt":
            return self._load_txt()
        if ext in (".docx", ".doc"):
            return self._load_docx()
        raise ValueError(f"Unsupported format: {ext}")

    # ------------------------------------------------------------------
    @staticmethod
    def _patch_ebooklib() -> None:
        """Monkey-patch ebooklib to tolerate EPUBs with broken NCX files."""
        import ebooklib.epub as _em
        _orig = _em.EpubReader._parse_ncx
        def _safe_parse_ncx(self, file_name):
            try:
                _orig(self, file_name)
            except (AttributeError, TypeError):
                pass
        _em.EpubReader._parse_ncx = _safe_parse_ncx

    # ------------------------------------------------------------------
    def _load_epub(self) -> tuple[list[dict], str]:
        """Read an EPUB and return (chapters, clean title)."""
        self._patch_ebooklib()
        book  = epub.read_epub(str(self.epub_path), options={"ignore_ncx": True})
        meta  = book.get_metadata("DC", "title")
        raw   = meta[0][0].strip() if meta else self.epub_path.stem
        title = self._clean_title(raw) or self.epub_path.stem

        items       = {item.get_id(): item for item in book.get_items()}
        chapters:   list[dict] = []
        skipped_toc = 0
        doc_idx     = -1   # sequential index among ITEM_DOCUMENT spine items

        for spine_id, _ in book.spine:
            item = items.get(spine_id)
            if item is None or item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue
            doc_idx += 1
            if doc_idx in self.skip_chapters:
                continue

            segments = html_to_segments(item.get_content())
            if not segments:
                continue

            # Merge consecutive paragraphs that don't end in a real sentence
            # boundary (e.g. each dialogue line is often its own <p> tag in
            # EPUB markup) so they share one edge-tts call/pause instead of
            # each getting a hard paragraph-splice pause.
            segments = merge_incomplete_segments(segments)

            plain = segments_to_plain_text(segments)
            if len(plain) < self.skip_short:
                continue

            # Determine chapter title from first heading or item name
            soup     = BeautifulSoup(item.get_content(), "html.parser")
            h_tag    = soup.find(["h1", "h2", "h3"])
            ch_title = h_tag.get_text(strip=True) if h_tag else item.get_name()

            # Skip TOC chapters (nav tags are already removed by html_to_segments;
            # this catches TOC chapters that appear in the spine)
            if is_toc_title(ch_title) or is_toc_content(plain):
                skipped_toc += 1
                print(f"  ⏭  Inhaltsverzeichnis übersprungen: {ch_title!r}")
                continue

            chapters.append({"title": ch_title, "segments": segments})

        if skipped_toc:
            print(f"  ℹ  {skipped_toc} Inhaltsverzeichnis-Kapitel übersprungen.\n")

        return chapters, title

    # ------------------------------------------------------------------
    def _load_pdf(self) -> tuple[list[dict], str]:
        """
        Read a PDF and return (chapters, clean title).
        Attempts heading detection via font-size comparison (Option 1).
        Falls back to paragraph-only structure with a warning if sizes are uniform.
        """
        import fitz  # pymupdf

        doc   = fitz.open(str(self.epub_path))
        raw   = (doc.metadata.get("title") or "").strip() or self.epub_path.stem
        title = self._clean_title(raw) or self.epub_path.stem

        # --- Detect heading font-size threshold ---
        font_sizes: list[float] = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:   # 0 = text block
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("text", "").strip():
                            font_sizes.append(span["size"])

        has_heading_detection = False
        heading_threshold     = float("inf")
        if font_sizes:
            sorted_sizes = sorted(font_sizes)
            median_size  = sorted_sizes[len(sorted_sizes) // 2]
            max_size     = sorted_sizes[-1]
            # Only use heading detection when there is meaningful size variation (≥15%)
            if max_size > median_size * 1.15:
                heading_threshold     = median_size * 1.15
                has_heading_detection = True

        if not has_heading_detection:
            print(
                "  ℹ  Hinweis (PDF): Überschriften konnten nicht per Schriftgröße erkannt werden –\n"
                "     alle Textblöcke werden als Absätze behandelt (SSML-Absatz-Pausen aktiv).\n"
                "     Tipp: Option 3 (Platzhalter-Sätze) wäre eine Alternative für\n"
                "     noch klarere Strukturierung bei uniform gesetzten PDFs.",
                file=sys.stderr,
            )

        # --- Apply user-selected page range (1-based, inclusive) ---
        first_page = max(1, self.start_page or 1)
        last_page  = min(len(doc), self.end_page or len(doc))

        # --- Extract segments page by page, group into chapters ---
        TARGET    = 15_000
        chapters: list[dict] = []
        cur_segs: list[dict] = []
        cur_len   = 0
        page_start = first_page
        skipped_toc = 0

        for i, page in enumerate(doc, start=1):
            if i < first_page or i > last_page:
                continue
            if has_heading_detection:
                for block in page.get_text("dict")["blocks"]:
                    if block.get("type") != 0:
                        continue
                    block_parts: list[str] = []
                    is_heading = False
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            t = sanitize_text(span.get("text", ""))
                            if t:
                                block_parts.append(t)
                                if span["size"] >= heading_threshold:
                                    is_heading = True
                    block_text = " ".join(block_parts).strip()
                    if block_text:
                        seg_type = SEG_HEADING if is_heading else SEG_PARAGRAPH
                        cur_segs.append({"type": seg_type, "text": block_text})
                        cur_len += len(block_text)
            else:
                page_text = sanitize_text(page.get_text())
                if page_text:
                    cur_segs.append({"type": SEG_PARAGRAPH, "text": page_text})
                    cur_len += len(page_text)

            at_end = (i == len(doc))
            if cur_len >= TARGET or at_end:
                plain = segments_to_plain_text(cur_segs)
                label = f"Seiten {page_start}–{i}" if page_start < i else f"Seite {i}"
                if len(plain) >= self.skip_short:
                    if is_toc_content(plain):
                        skipped_toc += 1
                        print(f"  ⏭  Inhaltsverzeichnis übersprungen: {label}")
                    else:
                        chapters.append({"title": label, "segments": cur_segs})
                cur_segs   = []
                cur_len    = 0
                page_start = i + 1

        doc.close()
        if skipped_toc:
            print(f"  ℹ  {skipped_toc} Inhaltsverzeichnis-Seite(n) übersprungen.\n")

        return chapters, title

    # ------------------------------------------------------------------
    def _load_txt(self) -> tuple[list[dict], str]:
        """
        Read a plain-text file and return (sections as chapters, clean title).
        Detects headings via common keywords; splits on double-newlines for paragraphs.
        """
        raw_title = self._clean_title(self.epub_path.stem) or self.epub_path.stem
        content   = self.epub_path.read_text(encoding="utf-8", errors="replace")

        # Remove mid-sentence/mid-dialogue line breaks so each paragraph is
        # a clean, uninterrupted unit of text (avoids unwanted TTS pauses).
        content = prepare_text_for_edge_tts(content)

        # Optionally save the preprocessed file next to the original.
        if self.translate_to:
            pre_path = self.epub_path.with_name(
                self.epub_path.stem + "_preprocessed.txt"
            )
            pre_path.write_text(content, encoding="utf-8")
            print(f"  📝  Vorverarbeitete Datei gespeichert: {pre_path.name}")

        HEADING_RE = re.compile(
            r"(?m)^(?P<kw>Kapitel|Chapter|Teil|Part|Abschnitt|Section|Epilog|Prolog|Vorwort|Nachwort)\s+\S.*$"
        )
        # "Abschnitt"/"Section" reads as a sub-heading within a chapter; everything
        # else (Teil/Part, Kapitel/Chapter, Epilog/Prolog/Vorwort/Nachwort) is level 1.
        TXT_HEADING_LEVEL = {"abschnitt": 2, "section": 2}

        # Build flat list of segments, marking headings
        segments_raw: list[dict] = []
        last_end = 0

        for m in HEADING_RE.finditer(content):
            # Paragraphs before this heading
            before = content[last_end:m.start()]
            for para in re.split(r"\n{2,}", before):
                text = sanitize_text(para)
                if len(text) >= self.skip_short:
                    segments_raw.append({"type": SEG_PARAGRAPH, "text": text})
            # The heading line
            heading_text = sanitize_text(m.group(0))
            if heading_text:
                level = TXT_HEADING_LEVEL.get(m.group("kw").lower(), 1)
                segments_raw.append({"type": SEG_HEADING, "text": heading_text, "level": level})
            last_end = m.end()

        # Remaining text after the last heading
        for para in re.split(r"\n{2,}", content[last_end:]):
            text = sanitize_text(para)
            if len(text) >= self.skip_short:
                segments_raw.append({"type": SEG_PARAGRAPH, "text": text})

        # No headings found → split on paragraph breaks only
        if not segments_raw:
            for para in re.split(r"\n{2,}", content):
                text = sanitize_text(para)
                if len(text) >= self.skip_short:
                    segments_raw.append({"type": SEG_PARAGRAPH, "text": text})

        # Further fallback: no paragraphs → equal-size chunks
        if not segments_raw:
            full  = sanitize_text(content)
            CHUNK = 20_000
            for start in range(0, len(full), CHUNK):
                text = full[start:start + CHUNK]
                if len(text) >= self.skip_short:
                    segments_raw.append({"type": SEG_PARAGRAPH, "text": text})

        # Merge consecutive paragraphs that end mid-sentence
        segments_raw = merge_incomplete_segments(segments_raw)

        # Group flat segments into logical chapters (split at heading boundaries)
        chapters:    list[dict] = []
        current:     list[dict] = []
        skipped_toc  = 0
        sec_idx      = -1   # sequential section index (for skip_chapters)

        def flush_chapter(segs: list[dict]) -> None:
            nonlocal skipped_toc, sec_idx
            if not segs:
                return
            sec_idx += 1
            if sec_idx in self.skip_chapters:
                return
            plain            = segments_to_plain_text(segs)
            title_candidate  = (
                segs[0]["text"][:80]
                if segs[0]["type"] == SEG_HEADING
                else f"Abschnitt {len(chapters) + 1}"
            )
            if is_toc_title(title_candidate) or is_toc_content(plain):
                skipped_toc += 1
                print(f"  ⏭  Inhaltsverzeichnis übersprungen: {title_candidate!r}")
            else:
                chapters.append({"title": title_candidate, "segments": segs})

        for seg in segments_raw:
            if seg["type"] == SEG_HEADING and current:
                flush_chapter(current)
                current = [seg]
            else:
                current.append(seg)
        flush_chapter(current)

        if skipped_toc:
            print(f"  ℹ  {skipped_toc} Inhaltsverzeichnis-Abschnitt(e) übersprungen.\n")

        return chapters, raw_title

    # ------------------------------------------------------------------
    def _load_docx(self) -> tuple[list[dict], str]:
        """
        Read a .docx (or .doc) file and return (chapters, clean title).
        Iterates all paragraphs in document order including table cells.
        Chapters are split at heading-style paragraphs; falls back to
        one chapter if no headings are found.
        For .doc files a conversion is attempted via pywin32 or LibreOffice.
        """
        from docx import Document  # python-docx

        src = self.epub_path
        tmp_dir: "Path | None" = None

        if src.suffix.lower() == ".doc":
            src, tmp_dir = _convert_doc_to_docx(src)

        try:
            doc = Document(str(src))
        finally:
            if tmp_dir is not None:
                import shutil
                shutil.rmtree(tmp_dir, ignore_errors=True)

        raw   = (doc.core_properties.title or "").strip()
        title = self._clean_title(raw) or self.epub_path.stem

        CHAPTER_STYLES = {"heading 1", "überschrift 1", "titre 1", "título 1"}
        HEADING_LEVEL_RE = re.compile(r"(\d+)\s*$")

        # Build flat segment list – walk body + table cells in document order
        segments_raw: list[dict] = []
        heading_styles_seen: set[str] = set()
        for para in _docx_iter_paragraphs(doc):
            text = sanitize_text(para.text)
            if not text:
                continue
            style_lower = para.style.name.lower() if para.style else ""
            is_hdg = "heading" in style_lower or "überschrift" in style_lower
            seg: dict = {"type": SEG_HEADING if is_hdg else SEG_PARAGRAPH, "text": text}
            if is_hdg:
                m = HEADING_LEVEL_RE.search(style_lower)
                seg["level"] = int(m.group(1)) if m else 1
                seg["_chapter_boundary_style"] = style_lower
                heading_styles_seen.add(style_lower)
            segments_raw.append(seg)

        if not segments_raw:
            return [], title

        # Merge consecutive paragraphs that end mid-sentence (no trailing punctuation)
        segments_raw = merge_incomplete_segments(segments_raw)

        # A chapter boundary is a "Heading 1"-equivalent style, so that H2+
        # sub-headings stay *inside* their chapter (audibly distinguished via
        # pitch/rate/pause instead of splitting into their own file). Falls
        # back to "any heading" for documents that don't use that convention
        # at all (e.g. only "Heading 2" is used throughout).
        chapter_boundary_styles = heading_styles_seen & CHAPTER_STYLES or heading_styles_seen

        # Group flat segments into logical chapters (split at chapter-boundary headings)
        chapters: list[dict] = []
        current:  list[dict] = []

        def _flush(segs: list[dict]) -> None:
            if not segs:
                return
            plain = segments_to_plain_text(segs)
            if len(plain) < self.skip_short:
                return
            ch_title = (
                segs[0]["text"][:80] if segs[0]["type"] == SEG_HEADING
                else f"Abschnitt {len(chapters) + 1}"
            )
            if is_toc_title(ch_title) or is_toc_content(plain):
                print(f"  ⏭  Inhaltsverzeichnis übersprungen: {ch_title!r}")
                return
            chapters.append({"title": ch_title, "segments": segs})

        for seg in segments_raw:
            is_boundary = seg["type"] == SEG_HEADING and (
                seg.pop("_chapter_boundary_style", None) in chapter_boundary_styles
            )
            if is_boundary and current:
                _flush(current)
                current = [seg]
            else:
                current.append(seg)
        _flush(current)

        # Fallback: no headings at all → one chapter with everything
        if not chapters:
            plain = segments_to_plain_text(segments_raw)
            if len(plain) >= self.skip_short:
                chapters.append({"title": title, "segments": segments_raw})

        return chapters, title

    # ------------------------------------------------------------------
    @staticmethod
    def _is_network_error(exc: Exception) -> bool:
        """Return True if the exception looks like a server/network connectivity failure."""
        msg = str(exc).lower()
        return any(k in msg for k in (
            "getaddrinfo", "cannot connect", "connection refused",
            "network is unreachable", "timed out", "ssl", "websocket",
            "service unavailable", "503", "502",
        ))

    # ------------------------------------------------------------------
    @staticmethod
    def _ffmpeg_path() -> str | None:
        """Return the path to the ffmpeg executable, or None if not installed."""
        return shutil.which("ffmpeg")

    @staticmethod
    def _normalize_mp3_loudness(path: Path, target_lufs: float = -16.0) -> bool:
        """
        In-place loudness-normalize an MP3 via ffmpeg's loudnorm filter
        (single-pass, EBU R128). Leaves the file untouched on failure.
        Returns True on success.
        """
        ffmpeg = EpubConverter._ffmpeg_path()
        if not ffmpeg:
            return False
        tmp = path.with_suffix(".norm.mp3")
        try:
            result = subprocess.run(
                [
                    ffmpeg, "-y", "-i", str(path),
                    "-af", f"loudnorm=I={target_lufs}:TP=-1.5:LRA=11",
                    "-ar", "44100", "-b:a", "128k",
                    str(tmp),
                ],
                capture_output=True,
            )
        except OSError:
            tmp.unlink(missing_ok=True)
            return False
        if result.returncode == 0 and tmp.exists() and tmp.stat().st_size > 0:
            # On Windows a virus scanner can briefly hold a lock on a just-written
            # file; retry the rename a few times before giving up.
            for attempt in range(5):
                try:
                    tmp.replace(path)
                    return True
                except OSError:
                    if attempt == 4:
                        tmp.unlink(missing_ok=True)
                        return False
                    time.sleep(0.2)
        tmp.unlink(missing_ok=True)
        return False

    @staticmethod
    def _concat_mp3_ffmpeg(parts: list[Path], merged: Path) -> bool:
        """
        Properly concatenate MP3 files via ffmpeg's concat demuxer
        (stream copy, no re-encoding). Returns True on success.
        """
        ffmpeg = EpubConverter._ffmpeg_path()
        if not ffmpeg:
            return False
        list_file = merged.with_suffix(".concat.txt")
        list_file.write_text(
            "\n".join(f"file '{p.resolve().as_posix()}'" for p in parts),
            encoding="utf-8",
        )
        try:
            result = subprocess.run(
                [
                    ffmpeg, "-y", "-f", "concat", "-safe", "0",
                    "-i", str(list_file), "-c", "copy", str(merged),
                ],
                capture_output=True,
            )
        except OSError:
            result = None
        list_file.unlink(missing_ok=True)
        return bool(result) and result.returncode == 0 and merged.exists() and merged.stat().st_size > 0

    # ------------------------------------------------------------------
    @staticmethod
    def _piper_model_dir() -> Path:
        """Return (and create) the directory where Piper models are stored."""
        import os
        if sys.platform == "win32":
            base = Path(os.environ.get("APPDATA", Path.home()))
        else:
            base = Path(os.environ.get("XDG_DATA_HOME", Path.home() / ".local" / "share"))
        d = base / "epub-to-voice" / "piper-models"
        d.mkdir(parents=True, exist_ok=True)
        return d

    @staticmethod
    def _ensure_piper_model(voice_name: str) -> tuple[Path, Path]:
        """
        Return (onnx_path, json_path) for a Piper voice.
        Downloads both files from HuggingFace if they are missing.
        """
        import urllib.request
        mdir      = EpubConverter._piper_model_dir()
        onnx_path = mdir / f"{voice_name}.onnx"
        json_path = mdir / f"{voice_name}.onnx.json"

        if not onnx_path.exists() or not json_path.exists():
            # Parse voice name:  de_DE-thorsten-high  →  de / de_DE / thorsten / high
            parts       = voice_name.split("-")
            lang_region = parts[0]                       # de_DE
            lang        = lang_region.split("_")[0].lower()  # de
            quality     = parts[-1]                      # high
            name        = "-".join(parts[1:-1])          # thorsten

            base_url = (
                f"https://huggingface.co/rhasspy/piper-voices"
                f"/resolve/v1.0.0/{lang}/{lang_region}/{name}/{quality}"
            )
            for fname, fpath in [
                (f"{voice_name}.onnx",      onnx_path),
                (f"{voice_name}.onnx.json", json_path),
            ]:
                if not fpath.exists():
                    url = f"{base_url}/{fname}"
                    print(f"  ⬇  Lade Piper-Modell: {fname} …", flush=True)
                    urllib.request.urlretrieve(url, str(fpath))
                    print(f"     ✓ {fname}")

        return onnx_path, json_path

    def _get_piper_voice(self):
        """Lazy-load and cache the PiperVoice instance."""
        if self._piper_instance is None:
            from piper.voice import PiperVoice  # type: ignore
            onnx_path, json_path = EpubConverter._ensure_piper_model(self.piper_voice)
            print(f"  🔊  Initialisiere Piper: {self.piper_voice} …", flush=True)
            self._piper_instance = PiperVoice.load(
                str(onnx_path), config_path=str(json_path), use_cuda=False
            )
        return self._piper_instance

    # ------------------------------------------------------------------
    async def _synthesise_chunk_piper_legacy(
        self, segments: list[dict], out_path: Path,
        log: "HtmlLog | None" = None,
    ) -> None:
        """
        Synthesise a chunk with Piper TTS (offline) in a single request.

        No paragraph/heading pause-splicing - used when ffmpeg is not
        installed or "Optimiere vor dem Lesen" is disabled. See
        _synthesise_chunk_piper_spliced() for the pause-aware path.
        """
        text = segments_to_plain_text(segments)
        if not text.strip():
            return
        await self._synthesise_piece_piper(text, self.rate, out_path)

    # ------------------------------------------------------------------
    async def _synthesise_piece_piper(self, text: str, rate: str, out_path: Path) -> int:
        """
        Synthesise one text piece via Piper TTS and encode to MP3.
        Returns the sample rate of the generated audio (voice-model dependent,
        needed so silence clips spliced in alongside it can match exactly).
        """
        import lameenc  # type: ignore
        from piper.config import SynthesisConfig  # type: ignore

        voice  = self._get_piper_voice()
        length = _rate_to_piper_length_scale(rate)
        cfg    = SynthesisConfig(length_scale=length)

        # synthesize() returns an iterable of AudioChunk objects (raw int16 PCM)
        chunks = list(voice.synthesize(text, syn_config=cfg))
        if not chunks:
            raise RuntimeError("Piper TTS lieferte keine Audio-Daten.")

        first   = chunks[0]
        pcm_all = b"".join(ch.audio_int16_bytes for ch in chunks)

        # Encode PCM → MP3 via lameenc (no ffmpeg required)
        enc = lameenc.Encoder()
        enc.set_bit_rate(128)
        enc.set_in_sample_rate(first.sample_rate)
        enc.set_channels(first.sample_channels)
        enc.set_quality(2)   # 2 = highest quality (LAME preset)
        mp3_data = enc.encode(pcm_all) + enc.flush()

        out_path.write_bytes(mp3_data)

        if out_path.stat().st_size < 512:
            out_path.unlink(missing_ok=True)
            raise RuntimeError("Piper TTS: MP3-Ausgabe ist leer.")

        return first.sample_rate

    # ------------------------------------------------------------------
    async def _synthesise_chunk_piper_spliced(
        self, segments: list[dict], out_path: Path,
        log: "HtmlLog | None" = None,
    ) -> None:
        """
        Synthesise a chunk with Piper TTS, one piece per paragraph/heading,
        spliced together with real silence gaps (ffmpeg) - the same approach
        _synthesise_chunk_edge() uses, since Piper (like edge-tts) has no
        reliable inline pause markup (see plan_audio_pieces()). Without this,
        Piper reads straight through chapter/paragraph/heading boundaries
        with no pause at all.
        """
        pieces = plan_audio_pieces(segments, self.rate)
        tts_pieces = [p for p in pieces if p["kind"] == "tts"]
        if not tts_pieces:
            return

        tmp_dir = out_path.parent / f".tmp_{out_path.stem}"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        try:
            # Two passes: Piper's sample rate (voice-model dependent) is only
            # known after the first synthesis call, but silence clips need to
            # match it exactly for ffmpeg's stream-copy concat to work.
            resolved: list[Path | None] = [None] * len(pieces)
            sample_rate = 22050  # Piper's common default; overwritten below
            for idx, piece in enumerate(pieces):
                if piece["kind"] != "tts":
                    continue
                part_path = tmp_dir / f"{idx:04d}.mp3"
                sample_rate = await self._synthesise_piece_piper(piece["text"], piece["rate"], part_path)
                resolved[idx] = part_path

            part_paths: list[Path] = []
            for idx, piece in enumerate(pieces):
                if piece["kind"] == "silence":
                    part_paths.append(self._get_silence_clip(piece["ms"], sample_rate=sample_rate))
                else:
                    part_paths.append(resolved[idx])

            if not self._concat_mp3_ffmpeg(part_paths, out_path):
                raise RuntimeError("Zusammenfügen der Segment-Audios via ffmpeg fehlgeschlagen.")
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    # ------------------------------------------------------------------
    async def _synthesise_chunk(
        self, segments: list[dict], out_path: Path,
        log: "HtmlLog | None" = None,
    ) -> None:
        """Dispatch to the right TTS engine (and pause-splicing vs. legacy mode)."""
        if self.tts_engine == "piper":
            if self.optimize_before_reading and self._ffmpeg_path():
                await self._synthesise_chunk_piper_spliced(segments, out_path, log=log)
            else:
                await self._synthesise_chunk_piper_legacy(segments, out_path, log=log)
        else:
            if self.optimize_before_reading:
                await self._synthesise_chunk_edge(segments, out_path, log=log)
            else:
                await self._synthesise_chunk_edge_legacy(segments, out_path, log=log)

    # ------------------------------------------------------------------
    async def _synthesise_chunk_edge_legacy(
        self, segments: list[dict], out_path: Path,
        log: "HtmlLog | None" = None,
    ) -> None:
        """
        Synthesise a segment chunk to MP3 via Edge TTS in a single request.
        Primary:  punctuation-based prosodic pauses (Option 1).
        Fallback: spoken placeholder sentences (Option 3) if synthesis fails.
        Retries up to 3 times on network/server errors with exponential backoff (2 s, 4 s).
        Raises TtsServerError only after all retries are exhausted.

        Used when ffmpeg is not installed (the pause-splicing path in
        _synthesise_chunk_edge requires ffmpeg to generate silence and
        stitch per-segment audio together).
        """
        MAX_RETRIES = 3
        BASE_DELAY  = 2.0   # seconds; doubles each attempt

        def _warn(msg: str) -> None:
            if log:
                log.log(msg, "warn")
            else:
                print(msg)

        for attempt in range(1, MAX_RETRIES + 1):
            last_network_exc: Exception | None = None

            # --- Primary: Option 1 ---
            ssml_text = segments_to_tts_text(segments)
            try:
                communicate = edge_tts.Communicate(ssml_text, self.voice, rate=self.rate, volume=self.volume)
                await communicate.save(str(out_path))
                if out_path.exists() and out_path.stat().st_size < 1024:
                    out_path.unlink()
                    raise RuntimeError("Edge-TTS lieferte eine leere/korrupte Datei.")
                return  # success
            except Exception as exc:
                if self._is_network_error(exc):
                    last_network_exc = exc
                else:
                    # Non-network error → try Option 3 fallback
                    _warn(f"  ⚠  Synthesis fehlgeschlagen ({exc}) – wechsle zu Option 3 …")
                    fallback = segments_to_fallback_text(segments)
                    try:
                        communicate = edge_tts.Communicate(fallback, self.voice, rate=self.rate, volume=self.volume)
                        await communicate.save(str(out_path))
                        if out_path.exists() and out_path.stat().st_size < 1024:
                            out_path.unlink()
                            raise RuntimeError("Leere Datei.")
                        return  # success with fallback
                    except Exception as exc2:
                        if self._is_network_error(exc2):
                            last_network_exc = exc2
                        else:
                            raise RuntimeError(
                                f"Edge-TTS: Auch Fallback-Text (Option 3) fehlgeschlagen: {exc2}"
                            ) from exc2

            # Network error reached – retry with backoff or give up
            if attempt < MAX_RETRIES:
                delay = BASE_DELAY * (2 ** (attempt - 1))   # 2 s, 4 s
                msg = f"  ↻  Server nicht erreichbar – Versuch {attempt}/{MAX_RETRIES}, warte {delay:.0f} s …"
                _warn(msg)
                print(msg, flush=True)
                await asyncio.sleep(delay)
            else:
                raise TtsServerError(str(last_network_exc)) from last_network_exc

    # ------------------------------------------------------------------
    @staticmethod
    def _get_silence_clip(ms: int, sample_rate: int = 24000) -> Path:
        """
        Return a cached silent MP3 of the given duration (mono, at sample_rate,
        matching the engine's own output so ffmpeg concat -c copy stays clean -
        Edge TTS uses 24 kHz; Piper voices commonly use 22.05 kHz).
        Generated once via ffmpeg and reused across all conversions.
        """
        cache_dir = Path(tempfile.gettempdir()) / "epub_to_voice_silence"
        cache_dir.mkdir(parents=True, exist_ok=True)
        path = cache_dir / f"silence_{ms}ms_{sample_rate}.mp3"
        if path.exists() and path.stat().st_size > 0:
            return path

        ffmpeg = EpubConverter._ffmpeg_path()
        subprocess.run(
            [
                ffmpeg, "-y", "-f", "lavfi",
                "-i", f"anullsrc=channel_layout=mono:sample_rate={sample_rate}",
                "-t", f"{max(ms, 1) / 1000:.3f}",
                "-c:a", "libmp3lame", "-b:a", "48k",
                str(path),
            ],
            capture_output=True,
        )
        return path

    # ------------------------------------------------------------------
    async def _synthesise_piece_edge(
        self, text: str, pitch: str, rate: str, out_path: Path,
        log: "HtmlLog | None" = None,
    ) -> None:
        """
        Synthesise one short text piece to MP3 via Edge TTS.
        Retries up to 3 times on network/server errors (2 s / 4 s backoff),
        then falls back once to a short placeholder sentence (Option 3)
        on non-network errors before giving up.
        """
        MAX_RETRIES = 3
        BASE_DELAY  = 2.0

        def _warn(msg: str) -> None:
            if log:
                log.log(msg, "warn")
            else:
                print(msg)

        current_text  = text
        used_fallback = False
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                communicate = edge_tts.Communicate(
                    current_text, self.voice, rate=rate, volume=self.volume, pitch=pitch,
                )
                await communicate.save(str(out_path))
                if out_path.exists() and out_path.stat().st_size < 512:
                    out_path.unlink()
                    raise RuntimeError("Edge-TTS lieferte eine leere/korrupte Datei.")
                return  # success
            except Exception as exc:
                if self._is_network_error(exc):
                    if attempt < MAX_RETRIES:
                        delay = BASE_DELAY * (2 ** (attempt - 1))
                        _warn(f"  ↻  Server nicht erreichbar – Versuch {attempt}/{MAX_RETRIES}, warte {delay:.0f} s …")
                        await asyncio.sleep(delay)
                        continue
                    raise TtsServerError(str(exc)) from exc
                if not used_fallback:
                    _warn(f"  ⚠  Segment-Synthese fehlgeschlagen ({exc}) – wechsle zu Platzhaltertext …")
                    current_text  = "Weiter."
                    used_fallback = True
                    continue
                raise RuntimeError(
                    f"Edge-TTS: Auch Platzhaltertext fehlgeschlagen: {exc}"
                ) from exc

    # ------------------------------------------------------------------
    async def _synthesise_chunk_edge(
        self, segments: list[dict], out_path: Path,
        log: "HtmlLog | None" = None,
    ) -> None:
        """
        Synthesise a segment chunk to MP3 via Edge TTS.

        Each segment (heading / paragraph / scene-break) is synthesised
        individually with a heading-level-aware pitch/rate, then spliced
        together with real silence gaps (ffmpeg) so headings/sub-headings
        are audibly distinguishable and pauses have exact, natural durations
        (edge-tts XML-escapes in-band SSML/break tags, so this is the only
        way to get controllable pauses - see plan_audio_pieces()).

        Falls back to a single combined request without per-level
        differentiation (_synthesise_chunk_edge_legacy) if ffmpeg is not
        installed.
        """
        if not self._ffmpeg_path():
            await self._synthesise_chunk_edge_legacy(segments, out_path, log=log)
            return

        pieces = plan_audio_pieces(segments, self.rate)
        tts_pieces = [p for p in pieces if p["kind"] == "tts"]
        if not tts_pieces:
            return

        tmp_dir = out_path.parent / f".tmp_{out_path.stem}"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        try:
            part_paths: list[Path] = []
            for idx, piece in enumerate(pieces):
                if piece["kind"] == "silence":
                    part_paths.append(self._get_silence_clip(piece["ms"]))
                else:
                    part_path = tmp_dir / f"{idx:04d}.mp3"
                    await self._synthesise_piece_edge(
                        piece["text"], piece["pitch"], piece["rate"], part_path, log=log,
                    )
                    part_paths.append(part_path)
                    await asyncio.sleep(0.15)   # be gentle with the service between calls

            if not self._concat_mp3_ffmpeg(part_paths, out_path):
                raise RuntimeError("Zusammenfügen der Segment-Audios via ffmpeg fehlgeschlagen.")
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    # ------------------------------------------------------------------
    async def convert(self) -> None:
        """Main entry point – convert all chapters to MP3 files."""
        # Preflight: check piper-tts + lameenc are installed when Piper engine is selected
        if self.tts_engine == "piper":
            missing = []
            try:
                from piper.voice import PiperVoice  # noqa: F401
            except ImportError:
                missing.append("piper-tts")
            try:
                import lameenc  # noqa: F401
            except ImportError:
                missing.append("lameenc")
            if missing:
                print(
                    f"❌  Fehlende Pakete für Piper TTS: {', '.join(missing)}\n"
                    "    Bitte ausführen:  pip install piper-tts lameenc\n"
                    "    Oder Edge TTS (online) in der GUI wählen.",
                    file=sys.stderr,
                )
                sys.exit(1)

        self.output_dir.mkdir(parents=True, exist_ok=True)
        chapters, book_title = self._load_chapters()

        safe_book = re.sub(r'[\\/:*?"<>|]', "_", book_title)[:80]

        if self.translate_to:
            # In resume mode, skip translation if all audio files already exist
            skip_translation = False
            if self.resume:
                check_chapters = chapters[:self.max_chapters] if self.max_chapters else chapters
                if check_chapters:
                    skip_translation = all(
                        (self.output_dir / f"{ch_idx:03d}_{sp_idx:03d}_{safe_book}{self._audio_ext}").exists()
                        and (self.output_dir / f"{ch_idx:03d}_{sp_idx:03d}_{safe_book}{self._audio_ext}").stat().st_size > 1024
                        for ch_idx, chapter in enumerate(check_chapters, start=1)
                        for sp_idx in range(1, len(chunk_segments(chapter["segments"])) + 1)
                    )
                if skip_translation:
                    print("⏭  Alle Splits vorhanden – Übersetzung und Bereinigung werden übersprungen.")

            if not skip_translation:
                try:
                    chapters = translate_chapters(chapters, self.translate_to)
                except ImportError:
                    print(
                        "  ⚠  deep_translator nicht installiert – Übersetzung übersprungen.\n"
                        "     Installieren mit: pip install deep_translator",
                        file=sys.stderr,
                    )
                else:
                    try:
                        pdf_path = save_translated_pdf(
                            chapters, self.output_dir, book_title, self.epub_path.stem
                        )
                        print(f"📄  Übersetzung als PDF gespeichert: {pdf_path.name}")
                    except Exception as exc:
                        print(f"  ⚠  PDF-Export fehlgeschlagen: {exc}", file=sys.stderr)
                    try:
                        txt_path = save_translated_txt(
                            chapters, self.output_dir, book_title, self.epub_path.stem
                        )
                        txt_path.unlink(missing_ok=True)
                    except Exception as exc:
                        print(f"  ⚠  TXT-Prüfung fehlgeschlagen: {exc}", file=sys.stderr)

        if self.max_chapters:
            chapters = chapters[:self.max_chapters]

        if not chapters:
            print("⚠  Keine lesbaren Kapitel gefunden – ist die Datei gültig?", file=sys.stderr)
            sys.exit(1)

        # Pre-compute all (chapter_idx, split_idx, segment_chunk) tuples
        jobs: list[tuple[int, int, list[dict]]] = []
        for ch_idx, chapter in enumerate(chapters, start=1):
            for sp_idx, seg_chunk in enumerate(chunk_segments(chapter["segments"]), start=1):
                jobs.append((ch_idx, sp_idx, seg_chunk))

        total = len(jobs)

        # --- Open HTML log ---
        log_path = self.output_dir / f"{safe_book}_protokoll.html"
        engine_label = (
            f"Piper ({self.piper_voice})" if self.tts_engine == "piper"
            else f"Edge ({self.voice})"
        )
        log = HtmlLog(log_path, book_title, {
            "Engine":  engine_label,
            "Rate":    self.rate,
            "Lautst.": self.volume,
            "Ausgabe": str(self.output_dir),
            "Splits":  str(total),
        }, enabled=self.save_log)

        file_type = self.epub_path.suffix.lstrip(".").upper()
        log.log(f"📖  Buch:      {book_title}", "info")
        log.log(f"📄  Eingabe:   {self.epub_path.name}  [{file_type}]", "dim")
        if self.tts_engine == "piper":
            log.log(f"🔊  Engine:    Piper TTS (offline)  |  Stimme: {self.piper_voice}  |  Rate: {self.rate}", "info")
        else:
            log.log(f"🔊  Engine:    Edge TTS  |  Stimme: {self.voice}  |  Rate: {self.rate}  |  Vol: {self.volume}", "info")
        log.log(f"📂  Ausgabe:   {self.output_dir}", "dim")
        log.log(f"📑  Kapitel:   {len(chapters)}  |  Splits: {total}", "info")
        log.log("", "sep")

        produced: list[Path] = []
        consecutive_errors = 0
        MAX_CONSECUTIVE = 3   # abort after this many consecutive server errors
        # Jobs that failed with a server error → retried once after the main loop
        failed_jobs: list[tuple[str, list[dict], Path]] = []

        if self.resume:
            log.log("▶  Fortsetzen – bereits vorhandene Splits werden übersprungen.", "info")

        for counter, (ch_idx, sp_idx, seg_chunk) in enumerate(jobs, start=1):
            filename = self.output_dir / f"{ch_idx:03d}_{sp_idx:03d}_{safe_book}{self._audio_ext}"
            chars    = sum(len(s["text"]) for s in seg_chunk)
            prefix   = f"  [{counter:>3}/{total}]  {ch_idx:03d}_{sp_idx:03d}  ({chars:,} Zeichen)"

            # Resume mode: skip splits that already have a valid audio file
            if self.resume and filename.exists() and filename.stat().st_size > 1024:
                produced.append(filename)
                log.log(f"{prefix} ⏭  übersprungen (vorhanden)", "dim")
                print(f"{prefix} ⏭")
                continue

            print(f"{prefix} … ", end="", flush=True)

            try:
                await self._synthesise_chunk(seg_chunk, filename, log=log)
                produced.append(filename)
                consecutive_errors = 0
                log.log(f"{prefix} ✓", "ok")
                await asyncio.sleep(0.5)   # kurze Pause → verhindert Rate-Limiting
            except TtsServerError as exc:
                consecutive_errors += 1
                msg = f"{prefix} ✗  SERVER NICHT ERREICHBAR: {exc}"
                log.log(msg, "error")
                failed_jobs.append((prefix, seg_chunk, filename))
                if consecutive_errors >= MAX_CONSECUTIVE:
                    abort = (
                        "❌  TTS-Server antwortet nicht (speech.platform.bing.com).\n"
                        "    Mögliche Ursachen:\n"
                        "      • Microsoft-Dienst temporär nicht verfügbar\n"
                        "      • Keine Internetverbindung\n"
                        "      • Firewall blockiert den Zugriff\n"
                        "    Konvertierung abgebrochen. Bitte später erneut versuchen."
                    )
                    log.log(abort, "error")
                    log.close(total, len(produced))
                    print(abort, file=sys.stderr)
                    sys.exit(2)
            except Exception as exc:
                consecutive_errors = 0
                log.log(f"{prefix} ✗  FEHLER: {exc}", "error")

        # --- Nachproduktion fehlgeschlagener Chunks ---
        if failed_jobs:
            wait_sec = 20
            retry_msg = (
                f"\n🔁  {len(failed_jobs)} Split(s) fehlgeschlagen – "
                f"warte {wait_sec} s und versuche erneut …"
            )
            log.log(retry_msg.strip(), "warn")
            print(retry_msg, flush=True)
            await asyncio.sleep(wait_sec)
            log.log("", "sep")

            for prefix, seg_chunk, filename in failed_jobs:
                print(f"{prefix} (Wiederholung) … ", end="", flush=True)
                try:
                    await self._synthesise_chunk(seg_chunk, filename, log=log)
                    produced.append(filename)
                    log.log(f"{prefix} ↺ ✓", "ok")
                    print("↺ ✓")
                    await asyncio.sleep(0.5)
                except TtsServerError as exc:
                    log.log(f"{prefix} ↺ ✗  SERVER NICHT ERREICHBAR: {exc}", "error")
                    print("↺ ✗")
                except Exception as exc:
                    log.log(f"{prefix} ↺ ✗  FEHLER: {exc}", "error")
                    print("↺ ✗")

            produced.sort()   # Wiederholungen in korrekte Reihenfolge einordnen

        log.log("", "sep")

        # --- Lautstärke-Normalisierung zwischen Kapiteln (ffmpeg, best effort) ---
        if self.normalize_volume and produced:
            if self._ffmpeg_path():
                msg = f"🔊  Normalisiere Lautstärke von {len(produced)} Datei(en) … "
                print(msg, end="", flush=True)
                ok_count = sum(1 for part in produced if self._normalize_mp3_loudness(part))
                log.log(f"{msg}✓ {ok_count}/{len(produced)} normalisiert.", "ok")
            else:
                log.log("⚠  ffmpeg nicht gefunden – Lautstärke-Normalisierung übersprungen.", "warn")

        if self.merge and produced:
            merged = self.output_dir / f"{safe_book}{self._audio_ext}"
            msg = f"🔗  Zusammenführen von {len(produced)} Datei(en) → {merged.name} … "
            print(msg, end="", flush=True)
            if not self._concat_mp3_ffmpeg(produced, merged):
                # Fallback: naive byte-concatenation (works for consecutive MP3 frames,
                # used when ffmpeg is not installed)
                with open(merged, "wb") as fout:
                    for part in produced:
                        fout.write(part.read_bytes())
            for part in produced:
                part.unlink(missing_ok=True)
            log.log(f"{msg}✓ Zusammengeführt.", "ok")

        done_msg = f"✅  Fertig! {len(produced)}/{total} Splits. Ausgabe: {self.output_dir}"
        log.log(done_msg, "ok")
        end_time = datetime.now()
        duration = str(end_time - log._start).split(".")[0]
        log.log(f"🕒  Ende:      {end_time.strftime('%d.%m.%Y %H:%M:%S')}  |  Laufzeit: {duration}", "dim")
        log.close(total, len(produced))
        print(f"\n✅  Fertig! Dateien gespeichert in: {self.output_dir}")


# ---------------------------------------------------------------------------
# Document structure extraction (for GUI preview)
# ---------------------------------------------------------------------------

def get_structure_cmd(file_path: str) -> None:
    """
    Extract document structure (chapters / pages) and print JSON to stdout.
    Used by the Electron front-end via --get-structure for the preview modal.
    """
    path = Path(file_path)
    if not path.exists():
        print(json.dumps({"error": f"File not found: {file_path}"}))
        return
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            _structure_pdf(path)
        elif ext == ".epub":
            _structure_epub(path)
        elif ext == ".txt":
            _structure_txt(path)
        elif ext in (".docx", ".doc"):
            _structure_docx(path)
        else:
            print(json.dumps({"error": f"Unsupported: {ext}"}))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}))


def _structure_pdf(path: Path, max_thumbs: int = 40) -> None:
    import fitz, base64 as _b64

    doc   = fitz.open(str(path))
    title = (doc.metadata.get("title") or "").strip() or path.stem
    total = len(doc)

    pages = []
    for i in range(min(max_thumbs, total)):
        page  = doc[i]
        rect  = page.rect
        scale = 120 / max(rect.width, 1)
        pix   = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        thumb = _b64.b64encode(pix.tobytes("png")).decode()
        plain = sanitize_text(page.get_text())
        pages.append({
            "page":         i + 1,
            "thumbnail":    thumb,
            "textPreview":  plain[:200],
            "looksLikeToc": is_toc_content(plain),
            "looksLikeCover": i == 0 and len(plain) < 300,
        })

    doc.close()
    print(json.dumps({
        "type":       "pdf",
        "title":      title,
        "totalPages": total,
        "thumbCount": len(pages),
        "pages":      pages,
    }), flush=True)


def _structure_epub(path: Path) -> None:
    EpubConverter._patch_ebooklib()
    book  = epub.read_epub(str(path), options={"ignore_ncx": True})
    meta  = book.get_metadata("DC", "title")
    raw   = meta[0][0].strip() if meta else path.stem
    title = EpubConverter._clean_title(raw) or path.stem

    items    = {item.get_id(): item for item in book.get_items()}
    chapters = []
    doc_idx  = -1

    for spine_id, _ in book.spine:
        item = items.get(spine_id)
        if item is None or item.get_type() != ebooklib.ITEM_DOCUMENT:
            continue
        doc_idx += 1
        segments = html_to_segments(item.get_content())
        plain    = segments_to_plain_text(segments)
        soup     = BeautifulSoup(item.get_content(), "html.parser")
        h_tag    = soup.find(["h1", "h2", "h3"])
        ch_title = h_tag.get_text(strip=True) if h_tag else item.get_name()
        chapters.append({
            "index":    doc_idx,
            "title":    ch_title,
            "chars":    len(plain),
            "isToc":    is_toc_title(ch_title) or is_toc_content(plain),
        })

    print(json.dumps({
        "type":     "epub",
        "title":    title,
        "chapters": chapters,
    }), flush=True)


def _structure_txt(path: Path) -> None:
    content = path.read_text(encoding="utf-8", errors="replace")
    title   = EpubConverter._clean_title(path.stem) or path.stem

    HEADING_RE = re.compile(
        r"(?m)^(?:Kapitel|Chapter|Teil|Part|Abschnitt|Section|Epilog|Prolog|Vorwort|Nachwort)\s+\S.*$"
    )
    parts   = re.split(r"\x00|\n{3,}", HEADING_RE.sub("\x00\\g<0>", content))
    sections = []
    for i, part in enumerate(parts):
        text = sanitize_text(part)
        if len(text) < 30:
            continue
        first_line = part.strip().splitlines()[0][:80] if part.strip() else f"Abschnitt {i+1}"
        sections.append({
            "index": i,
            "title": sanitize_text(first_line),
            "chars": len(text),
            "isToc": is_toc_title(first_line) or is_toc_content(text),
        })

    print(json.dumps({
        "type":     "txt",
        "title":    title,
        "sections": sections,
    }), flush=True)


def _structure_docx(path: Path) -> None:
    from docx import Document  # python-docx
    doc   = Document(str(path))
    raw   = (doc.core_properties.title or "").strip()
    title = EpubConverter._clean_title(raw) or path.stem

    CHAPTER_STYLES = {"heading 1", "überschrift 1", "titre 1", "título 1"}
    sections = []
    sec_idx  = 0
    cur_title = title
    cur_chars = 0
    cur_is_toc = False

    for para in doc.paragraphs:
        text  = sanitize_text(para.text)
        style = para.style.name.lower() if para.style else ""
        if style in CHAPTER_STYLES:
            if cur_chars > 0:
                sections.append({"index": sec_idx, "title": cur_title,
                                  "chars": cur_chars, "isToc": cur_is_toc})
                sec_idx += 1
            cur_title  = text[:80] or f"Abschnitt {sec_idx + 1}"
            cur_chars  = len(text)
            cur_is_toc = is_toc_title(cur_title)
        else:
            cur_chars  += len(text)
            cur_is_toc  = cur_is_toc or is_toc_content(text)

    if cur_chars > 0:
        sections.append({"index": sec_idx, "title": cur_title,
                          "chars": cur_chars, "isToc": cur_is_toc})

    print(json.dumps({
        "type":     "docx",
        "title":    title,
        "sections": sections,
    }), flush=True)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _docx_iter_paragraphs(doc: "object"):
    """
    Yield every paragraph in a python-docx Document in document order.
    Uses iter() to recurse into ALL nested elements – this covers normal
    body paragraphs, table cells, AND WordprocessingShape text boxes
    (wps:txbx) which doc.paragraphs silently skips.
    """
    from docx.text.paragraph import Paragraph

    W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    for p_el in doc.element.body.iter(W_P):
        yield Paragraph(p_el, doc)


# ---------------------------------------------------------------------------
# DOC → DOCX conversion helper
# ---------------------------------------------------------------------------

def _convert_doc_to_docx(doc_path: Path) -> "tuple[Path, Path]":
    """
    Convert a legacy .doc file to .docx and return (docx_path, tmp_dir).
    The caller is responsible for deleting tmp_dir after use.
    Tries pywin32 Word COM first, then LibreOffice headless.
    Raises RuntimeError if neither is available.
    """
    import tempfile
    tmp_dir = Path(tempfile.mkdtemp())
    out     = tmp_dir / (doc_path.stem + ".docx")

    # --- pywin32 Word COM (Windows, requires Microsoft Word) ---
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            wb = word.Documents.Open(str(doc_path.resolve()))
            wb.SaveAs2(str(out.resolve()), FileFormat=16)  # 16 = wdFormatXMLDocument
            wb.Close(False)
        finally:
            word.Quit()
        if out.exists():
            return out, tmp_dir
    except ImportError:
        pass
    except Exception as exc:
        print(f"  ⚠  Word COM fehlgeschlagen ({exc}) – versuche LibreOffice …", file=sys.stderr)

    # --- LibreOffice headless ---
    import subprocess
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx",
             "--outdir", str(tmp_dir), str(doc_path)],
            capture_output=True, timeout=120,
        )
        if result.returncode == 0 and out.exists():
            return out, tmp_dir
    except FileNotFoundError:
        pass
    except subprocess.TimeoutExpired:
        print("  ⚠  LibreOffice-Konvertierung abgelaufen.", file=sys.stderr)

    raise RuntimeError(
        f"Kann '{doc_path.name}' nicht öffnen.\n"
        "  Bitte die Datei in Word oder LibreOffice als .docx speichern,\n"
        "  oder pywin32 installieren: pip install pywin32"
    )


# ---------------------------------------------------------------------------
# PDF export of translated text
# ---------------------------------------------------------------------------

def save_translated_pdf(
    chapters: list[dict],
    output_dir: Path,
    book_title: str,
    file_stem: str,
) -> Path:
    """Render translated chapters as a PDF and save it in output_dir."""
    import fitz

    safe_stem = re.sub(r'[\\/:*?"<>|]', "_", file_stem)[:80]
    pdf_path  = output_dir / f"{safe_stem}_uebersetzt.pdf"

    # Build minimal HTML representing the translated content
    html_parts = [
        "<!DOCTYPE html><html><body>",
        f"<h1>{_html.escape(book_title)}</h1>",
    ]
    for chapter in chapters:
        title = chapter.get("title", "")
        if title:
            html_parts.append(f"<h2>{_html.escape(title)}</h2>")
        for seg in chapter.get("segments", []):
            text = seg.get("text", "")
            if not text:
                continue
            if seg.get("type") == SEG_HEADING:
                html_parts.append(f"<h3>{_html.escape(text)}</h3>")
            else:
                html_parts.append(f"<p>{_html.escape(text)}</p>")
    html_parts.append("</body></html>")
    html = "\n".join(html_parts)

    # Use PyMuPDF Story for automatic text flow across pages
    mediabox = fitz.paper_rect("a4")
    margin   = 50
    where    = mediabox + (margin, margin, -margin, -margin)

    story  = fitz.Story(html=html)
    writer = fitz.DocumentWriter(str(pdf_path))

    more = True
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()

    writer.close()
    return pdf_path


def save_translated_txt(
    chapters: list[dict],
    output_dir: Path,
    book_title: str,
    file_stem: str,
) -> Path:
    """
    Save translated chapters as a plain-text file in output_dir.
    Each segment becomes one paragraph, headings are separated by blank lines.
    Also checks for leftover newlines in segment texts and warns about them.
    """
    safe_stem = re.sub(r'[\\/:*?"<>|]', "_", file_stem)[:80]
    txt_path  = output_dir / f"{safe_stem}_uebersetzt.txt"

    lines: list[str] = [book_title, ""]
    bad_segs = 0
    for chapter in chapters:
        segs = chapter.get("segments", [])
        title = chapter.get("title", "")
        # Only output the chapter title separately when the first segment is NOT
        # already a heading – otherwise the (translated) heading segment is output
        # below and the untranslated title would create a duplicate.
        first_is_heading = bool(segs and segs[0].get("type") == SEG_HEADING)
        if title and not first_is_heading:
            lines.append(title)
            lines.append("")
        for seg in segs:
            text = seg.get("text", "")
            if not text:
                continue
            if "\n" in text:
                bad_segs += 1
                text = sanitize_text(text)   # fix leftover translator newlines
                seg["text"] = text
            lines.append(text)
            lines.append("")

    txt_path.write_text("\n".join(lines), encoding="utf-8")

    if bad_segs:
        print(
            f"  ⚠  {bad_segs} Segment(e) enthielten unerwünschte Zeilenumbrüche "
            f"nach der Übersetzung – wurden bereinigt."
        )
    else:
        print("  ✔  Keine unerwünschten Zeilenumbrüche durch Übersetzung eingeführt.")

    return txt_path


# ---------------------------------------------------------------------------
# Translation
# ---------------------------------------------------------------------------

def _translate_texts(texts: list[str], target_lang: str, source_lang: str = "auto") -> list[str]:
    """
    Translate a list of texts using deep_translator (Google Translate).
    Batches texts into chunks to stay under the ~5 000-char API limit.
    Returns a list of translated strings aligned to the input list.
    """
    from deep_translator import GoogleTranslator

    SEP   = "\n[|||]\n"   # separator that Google Translate reliably preserves
    LIMIT = 4500           # safe per-request limit

    results = list(texts)  # pre-fill with originals as fallback

    batch_idx:  list[int] = []
    batch_parts: list[str] = []
    batch_len   = 0

    def _flush(indices: list[int], parts: list[str]) -> None:
        if not parts:
            return
        joined     = SEP.join(parts)
        translator = GoogleTranslator(source=source_lang, target=target_lang)
        translated = translator.translate(joined) or ""
        split      = translated.split(SEP)
        for i, idx in enumerate(indices):
            raw = split[i] if i < len(split) else parts[i]
            results[idx] = sanitize_text(raw) or parts[i]

    for idx, text in enumerate(texts):
        needed = len(text) + len(SEP)
        if batch_len + needed > LIMIT and batch_parts:
            _flush(batch_idx, batch_parts)
            batch_idx, batch_parts, batch_len = [], [], 0
        batch_idx.append(idx)
        batch_parts.append(text)
        batch_len += needed

    _flush(batch_idx, batch_parts)
    return results


def translate_chapters(chapters: list[dict], target_lang: str) -> list[dict]:
    """Translate all segment texts in all chapters to target_lang in-place."""
    all_texts: list[str]          = []
    seg_map:   list[tuple[int, int]] = []  # (chapter_idx, seg_idx)

    for ci, chapter in enumerate(chapters):
        for si, seg in enumerate(chapter["segments"]):
            all_texts.append(seg["text"])
            seg_map.append((ci, si))

    if not all_texts:
        return chapters

    print(f"🌐 Übersetze {len(all_texts)} Textblöcke nach {target_lang} …", flush=True)
    translated = _translate_texts(all_texts, target_lang)

    # Deep-copy chapters and write translations back
    import copy
    result = copy.deepcopy(chapters)
    for i, (ci, si) in enumerate(seg_map):
        result[ci]["segments"][si]["text"] = translated[i] or chapters[ci]["segments"][si]["text"]

    return result


# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------

def _extract_sample(path: Path, max_chars: int = 6000) -> str:
    """Extract a plain-text sample from EPUB / PDF / TXT for language detection."""
    ext = path.suffix.lower()
    try:
        if ext == ".txt":
            return sanitize_text(
                path.read_text(encoding="utf-8", errors="replace")
            )[:max_chars]

        if ext == ".epub":
            from ebooklib import epub as _epub
            book = _epub.read_epub(str(path), options={"ignore_ncx": True})
            buf = ""
            for item in book.get_items():
                if item.get_type() != ebooklib.ITEM_DOCUMENT:
                    continue
                soup = BeautifulSoup(item.get_content(), "html.parser")
                for tag in soup(["nav", "script", "style"]):
                    tag.decompose()
                buf += " " + soup.get_text(" ", strip=True)
                if len(buf) >= max_chars:
                    break
            return sanitize_text(buf)[:max_chars]

        if ext == ".pdf":
            import fitz
            doc = fitz.open(str(path))
            buf = ""
            for page in doc:
                buf += page.get_text()
                if len(buf) >= max_chars:
                    break
            doc.close()
            return sanitize_text(buf)[:max_chars]

        if ext in (".docx", ".doc"):
            from docx import Document  # python-docx
            src     = path
            tmp_dir = None
            if ext == ".doc":
                src, tmp_dir = _convert_doc_to_docx(path)
            try:
                doc = Document(str(src))
                buf = " ".join(
                    sanitize_text(p.text) for p in doc.paragraphs if p.text.strip()
                )
            finally:
                if tmp_dir is not None:
                    import shutil
                    shutil.rmtree(tmp_dir, ignore_errors=True)
            return buf[:max_chars]

    except Exception:
        pass
    return ""


def _detect_language_heuristic(text: str) -> tuple[str, float]:
    """
    Lightweight heuristic language detection based on character sets and
    common function words. Returns (language_code, confidence 0–1).
    """
    sample = text[:4000].lower()
    words  = set(re.findall(r"\b[a-zäöüßàâçéèêëîïôùûüœæñ]{2,}\b", sample))

    LANG_WORDS: dict[str, list[str]] = {
        "de": ["und", "der", "die", "das", "ist", "nicht", "mit", "ich", "auf", "auch", "ein", "eine", "des", "dem", "den"],
        "en": ["the", "and", "is", "are", "was", "for", "this", "that", "with", "have", "not", "from", "they", "which"],
        "fr": ["les", "des", "une", "est", "que", "qui", "dans", "pour", "pas", "par", "sur", "avec", "mais", "ont"],
        "es": ["los", "las", "una", "que", "con", "por", "para", "como", "más", "pero", "del", "son", "sus", "este"],
        "it": ["che", "non", "una", "della", "sono", "come", "per", "con", "del", "alla", "nel", "questo", "più"],
        "nl": ["van", "het", "een", "dat", "zijn", "ook", "niet", "voor", "met", "door", "maar", "bij", "wordt"],
        "pt": ["que", "não", "uma", "com", "por", "para", "mais", "dos", "nas", "seu", "sua", "mas", "pelo"],
        "ru": [],  # Cyrillic → detected separately
        "zh": [],  # CJK → detected separately
        "ja": [],  # CJK / Kana → detected separately
    }

    # CJK / Cyrillic shortcuts
    if re.search(r"[\u4e00-\u9fff]", text[:2000]):
        # Distinguish Japanese (Hiragana/Katakana) from Chinese
        if re.search(r"[\u3040-\u30ff]", text[:2000]):
            return "ja", 0.85
        return "zh", 0.85
    if re.search(r"[\u0400-\u04ff]", text[:2000]):
        return "ru", 0.85

    scores: dict[str, int] = {}
    for lang, kw in LANG_WORDS.items():
        if not kw:
            continue
        scores[lang] = sum(1 for w in kw if w in words)

    if not scores or max(scores.values()) == 0:
        return "en", 0.3   # safe default

    best    = max(scores, key=scores.get)
    total   = sum(scores.values())
    conf    = round(scores[best] / total, 2) if total else 0.3
    return best, min(conf, 0.95)


def detect_language_cmd(file_path: str) -> None:
    """
    Detect the language of a file and print JSON to stdout.
    Used by the Electron front-end via --detect-language.
    Output: {"language": "de", "confidence": 0.9, "method": "langdetect"|"heuristic"|"error"}
    """
    path = Path(file_path)
    if not path.exists():
        print(json.dumps({"language": None, "error": f"File not found: {file_path}"}))
        return

    sample = _extract_sample(path)
    if not sample or len(sample) < 30:
        print(json.dumps({"language": None, "error": "Not enough text to detect language"}))
        return

    # Try langdetect first (more accurate)
    try:
        from langdetect import detect_langs
        results = detect_langs(sample)
        if results:
            best     = results[0]
            lang     = best.lang.split("-")[0]   # "zh-cn" → "zh"
            conf     = round(best.prob, 2)
            print(json.dumps({"language": lang, "confidence": conf, "method": "langdetect"}))
            return
    except ImportError:
        pass   # langdetect not installed → fall through to heuristic
    except Exception:
        pass

    # Heuristic fallback
    lang, conf = _detect_language_heuristic(sample)
    print(json.dumps({"language": lang, "confidence": conf, "method": "heuristic"}))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

async def list_voices_cmd(filter_str: str = "") -> None:
    voices   = await edge_tts.list_voices()
    filtered = [v for v in voices if filter_str.lower() in v["ShortName"].lower()] if filter_str else voices
    print(f"{'Short Name':<35}  {'Locale':<10}  Gender")
    print("-" * 60)
    for v in filtered:
        print(f"{v['ShortName']:<35}  {v['Locale']:<10}  {v['Gender']}")
    print(f"\n{len(filtered)} voice(s) found.")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Lightweight EPUB / PDF / TXT → MP3 converter powered by edge-tts.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python epub_to_voice.py book.epub
  python epub_to_voice.py book.epub -o D:/Audiobooks/MyBook -v en-US-AriaNeural -r "+15%"
  python epub_to_voice.py --list-voices de
        """,
    )
    parser.add_argument("epub",            nargs="?", help="Path to the .epub, .pdf, .txt, .docx, or .doc file")
    parser.add_argument("-o", "--output",  help="Output directory (default: same folder as the input file)")
    parser.add_argument("-v", "--voice",   default=DEFAULT_VOICE,
                        help=f'TTS voice (default: {DEFAULT_VOICE}). Alternatives e.g. '
                             f'"de-DE-KatjaNeural"/"de-DE-SeraphinaMultilingualNeural" (female), '
                             f'"en-US-JennyNeural" (female), "en-US-GuyNeural" (male). '
                             f'Use --list-voices for the full catalog.')
    parser.add_argument("-r", "--rate",    default=DEFAULT_RATE,  help='Speech rate, e.g. "+15%%" or "-10%%"')
    parser.add_argument("--volume",        default=DEFAULT_VOLUME, help='Volume, e.g. "+10%%"')
    parser.add_argument("--skip-short",    type=int, default=60,
                        help="Skip chapters/sections with fewer than N characters (default: 60)")
    parser.add_argument("--max-chapters",  type=int, default=None,
                        help="Convert only the first N chapters (e.g. 3 for a preview)")
    parser.add_argument("--merge",         action="store_true",
                        help="Concatenate all splits into a single MP3 file")
    parser.add_argument("--list-voices",      nargs="?", const="", metavar="FILTER",
                        help='List available voices, optionally filtered (e.g. --list-voices de)')
    parser.add_argument("--detect-language",  action="store_true",
                        help='Detect the language of the input file and print JSON (used by GUI)')
    parser.add_argument("--get-structure",    action="store_true",
                        help='Extract document structure (chapters/pages) and print JSON (used by GUI)')
    parser.add_argument("--start-page",  type=int, default=None,
                        help='PDF: first page to convert (1-based, default: 1)')
    parser.add_argument("--end-page",    type=int, default=None,
                        help='PDF: last page to convert (1-based, inclusive, default: last)')
    parser.add_argument("--skip-chapters", default="",
                        help='Comma-separated doc-order indices of chapters/sections to skip')
    parser.add_argument("--translate-to", default="",
                        help='Translate text to this language before TTS, e.g. "de" for German (requires deep_translator)')
    parser.add_argument("--resume", action="store_true",
                        help='Resume interrupted conversion: skip splits whose audio file already exists')
    parser.add_argument("--tts-engine", default="edge", choices=["edge", "piper"],
                        help='TTS engine: "edge" (online, Microsoft) or "piper" (offline, local)')
    parser.add_argument("--piper-voice", default="de_DE-thorsten-high",
                        help='Piper voice name, e.g. "de_DE-thorsten-high" or "en_US-amy-medium"')
    parser.add_argument("--no-normalize-volume", action="store_false", dest="normalize_volume",
                        default=True,
                        help='Disable loudness normalization between chapters (requires ffmpeg; on by default)')
    parser.add_argument("--save-log", action="store_true",
                        help='Save the HTML Protokoll file to the output directory (off by default)')
    parser.add_argument("--no-optimize-before-reading", action="store_false",
                        dest="optimize_before_reading", default=True,
                        help='Disable pause-splicing at paragraph/heading/scene breaks before reading '
                             '(on by default; especially important for Piper TTS, which has no natural pauses otherwise)')
    return parser


def main() -> None:
    # On Windows without a console (e.g. spawned from Electron), the default
    # ProactorEventLoop can fail with aiohttp/edge-tts. SelectorEventLoop is safer.
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

    parser = build_parser()
    args   = parser.parse_args()

    if args.list_voices is not None:
        asyncio.run(list_voices_cmd(args.list_voices))
        return

    if args.detect_language:
        if not args.epub:
            print(json.dumps({"language": None, "error": "No file specified"}))
            sys.exit(1)
        detect_language_cmd(args.epub)
        return

    if args.get_structure:
        if not args.epub:
            print(json.dumps({"error": "No file specified"}))
            sys.exit(1)
        get_structure_cmd(args.epub)
        return

    if not args.epub:
        parser.print_help()
        sys.exit(1)

    SUPPORTED = {".epub", ".pdf", ".txt", ".docx", ".doc"}
    epub_path = Path(args.epub)
    if not epub_path.exists():
        print(f"Fehler: Datei nicht gefunden: {epub_path}", file=sys.stderr)
        sys.exit(1)
    if epub_path.suffix.lower() not in SUPPORTED:
        print(f"Fehler: Nicht unterstütztes Format '{epub_path.suffix}'. Unterstützt: {', '.join(sorted(SUPPORTED))}", file=sys.stderr)
        sys.exit(1)

    skip_set = {int(x) for x in args.skip_chapters.split(",") if x.strip().isdigit()}

    converter = EpubConverter(
        epub_path     = epub_path,
        output_dir    = args.output,
        voice         = args.voice,
        rate          = args.rate,
        volume        = args.volume,
        skip_short    = args.skip_short,
        max_chapters  = args.max_chapters,
        merge         = args.merge,
        start_page    = args.start_page,
        end_page      = args.end_page,
        skip_chapters = skip_set,
        translate_to  = args.translate_to or None,
        resume        = args.resume,
        tts_engine    = args.tts_engine,
        piper_voice   = args.piper_voice,
        normalize_volume = args.normalize_volume,
        save_log      = args.save_log,
        optimize_before_reading = args.optimize_before_reading,
    )
    asyncio.run(converter.convert())


if __name__ == "__main__":
    main()
